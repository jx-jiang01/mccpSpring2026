#!/usr/bin/env python3
"""
Improve markdown formatting by re-processing Word documents with better structure extraction.
Uses LLM to format based on actual Word document structure.
"""

import os
import sys
from docx import Document
from pathlib import Path
import requests
import re
from datetime import datetime
import time

def get_poe_api_key():
    """Read Poe API key from LLM/poe.md"""
    poe_file = Path(__file__).parent.parent / "LLM" / "poe.md"
    if poe_file.exists():
        with open(poe_file, 'r') as f:
            lines = [line.strip() for line in f.readlines() if line.strip() and not line.strip().startswith('http')]
            if lines:
                return lines[0]
    return None

def extract_structured_text_from_word(docx_path):
    """Extract text from Word document with proper structure detection."""
    doc = Document(docx_path)
    structured_content = []
    
    print(f"Extracting structured content from: {Path(docx_path).name}")
    print(f"  Total paragraphs: {len(doc.paragraphs)}")
    print(f"  Total tables: {len(doc.tables)}")
    
    # Process paragraphs with structure
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else "Normal"
        
        # Detect structure
        if style_name.startswith('Heading'):
            level = 1
            if 'Heading' in style_name:
                try:
                    level = int(style_name.split()[-1])
                except:
                    level = 1
            structured_content.append({
                'type': 'heading',
                'level': level,
                'text': text
            })
        elif style_name in ['List Paragraph', 'List Bullet']:
            structured_content.append({
                'type': 'bullet',
                'text': text
            })
        elif style_name == 'List Number':
            structured_content.append({
                'type': 'numbered',
                'text': text
            })
        elif 'Task' in text or 'Warm-up' in text:
            structured_content.append({
                'type': 'task',
                'text': text
            })
        else:
            structured_content.append({
                'type': 'paragraph',
                'text': text
            })
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        structured_content.append({
            'type': 'table',
            'data': table_data
        })
    
    return structured_content

def format_table_markdown(table_data):
    """Convert table data to markdown table format."""
    if not table_data or not table_data[0]:
        return ""
    
    md_lines = []
    
    # Header row
    header = table_data[0]
    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("|" + "|".join(["---"] * len(header)) + "|")
    
    # Data rows
    for row in table_data[1:]:
        if row and any(cell.strip() for cell in row):
            # Pad row if needed
            while len(row) < len(header):
                row.append("")
            md_lines.append("| " + " | ".join(row[:len(header)]) + " |")
    
    return "\n".join(md_lines)

def build_text_for_llm(structured_content):
    """Build text representation for LLM processing."""
    text_parts = []
    
    for item in structured_content:
        if item['type'] == 'heading':
            level = item['level']
            text_parts.append(f"\n{'#' * (level + 1)} {item['text']}\n")
        elif item['type'] == 'bullet':
            text_parts.append(f"- {item['text']}\n")
        elif item['type'] == 'numbered':
            text_parts.append(f"1. {item['text']}\n")
        elif item['type'] == 'task':
            text_parts.append(f"\n**{item['text']}**\n")
        elif item['type'] == 'table':
            text_parts.append(f"\n{format_table_markdown(item['data'])}\n")
        else:
            text_parts.append(f"{item['text']}\n")
    
    return "".join(text_parts)

def call_poe_api(prompt, api_key):
    """Call Poe API to format content."""
    if not api_key:
        return None
    
    try:
        url = "https://api.poe.com/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "Claude-3.5-Sonnet",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 20000,
            "temperature": 0.3
        }
        
        print("  Calling Poe API for formatting...")
        response = requests.post(url, headers=headers, json=payload, timeout=180)
        
        if response.status_code == 200:
            result = response.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            print(f"  ✓ API call successful ({len(content)} chars returned)")
            return content
        else:
            print(f"  API error: {response.status_code}")
            return None
    except Exception as e:
        print(f"  Error: {e}")
        return None

def improve_markdown_formatting(docx_path, output_md_path, api_key):
    """Improve markdown formatting based on Word document structure."""
    print(f"\n{'='*70}")
    print(f"Improving markdown: {Path(docx_path).name}")
    print(f"{'='*70}")
    
    # Extract structured content
    structured_content = extract_structured_text_from_word(docx_path)
    
    # Build text representation
    raw_text = build_text_for_llm(structured_content)
    print(f"  Extracted {len(structured_content)} structured elements")
    print(f"  Total text length: {len(raw_text)} characters")
    
    # Extract session number
    session_match = re.search(r'Session (\d+)', Path(docx_path).stem)
    session_num = session_match.group(1) if session_match else "X"
    
    # Process in sections if content is too long
    if len(raw_text) > 15000:
        print(f"  Content too long ({len(raw_text)} chars), processing in sections...")
        
        # Split into logical sections
        sections = []
        current_section = []
        
        for item in structured_content:
            if item['type'] == 'heading' and item['level'] <= 2:
                if current_section:
                    sections.append(current_section)
                current_section = [item]
            else:
                current_section.append(item)
        if current_section:
            sections.append(current_section)
        
        print(f"  Split into {len(sections)} sections")
        
        # Format each section
        formatted_parts = []
        formatted_parts.append(f"# MCCP6020 ADVANCED ENGLISH FOR ACADEMIC PURPOSES SESSION {session_num}\n\n---\n")
        
        for section_idx, section in enumerate(sections, 1):
            section_text = build_text_for_llm(section)
            
            prompt = f"""Format this section of a Word document into markdown. CRITICAL: Preserve ALL content exactly.

Requirements:
- Use proper markdown headings (##, ###, ####)
- Format lists: "- " for bullets, "1. " for numbered
- Make task titles bold: **Task X**
- Format URLs: [text](url)
- Format tables as markdown tables
- Preserve ALL text - every word must be included
- Do NOT summarize or omit anything

Section Content:
{section_text[:8000]}

Return ONLY the formatted markdown for this section:"""
            
            formatted_section = call_poe_api(prompt, api_key)
            if formatted_section:
                # Clean up
                formatted_section = re.sub(r'^[^\#]*', '', formatted_section, flags=re.MULTILINE)
                formatted_parts.append(formatted_section.strip())
                formatted_parts.append("\n")
            else:
                # Fallback: basic formatting
                formatted_parts.append(section_text)
                formatted_parts.append("\n")
        
        formatted = "\n".join(formatted_parts)
    else:
        # Process entire document
        prompt = f"""Format this Word document content into well-structured markdown. CRITICAL: Preserve ALL content exactly - do NOT summarize or omit anything.

Requirements:
1. Add proper title at top: "# MCCP6020 ADVANCED ENGLISH FOR ACADEMIC PURPOSES SESSION {session_num}"
2. Use proper markdown headings based on structure:
   - Main sections: ## Section Title
   - Subsections: ### Subsection Title
   - Sub-subsections: #### Title
3. Format lists properly:
   - Bullet points: "- item"
   - Numbered lists: "1. item"
4. Make task titles bold: **Task X**, **Warm-up Task**
5. Format tables as proper markdown tables
6. Format URLs as [text](url)
7. Add horizontal rules (---) between major sections
8. Preserve all text, examples, citations exactly
9. Do NOT add summaries or "[Continues...]"
10. Keep proper spacing between sections

Word Document Content:
{raw_text}

Return the complete formatted markdown with proper structure and ALL content preserved:"""

        formatted = call_poe_api(prompt, api_key)

    # Call LLM
    formatted = call_poe_api(prompt, api_key)
    
    if formatted:
        # Clean up any LLM prefixes
        formatted = re.sub(r'^[^\#]*', '', formatted, flags=re.MULTILINE)
        formatted = formatted.strip()
        
        # Add images section if needed
        images_dir = Path(output_md_path).parent / "images" / f"Session {session_num}_Handout_word"
        if images_dir.exists():
            image_files = sorted(list(images_dir.glob("*.png")) + list(images_dir.glob("*.jpeg")) + list(images_dir.glob("*.jpg")))
            if image_files:
                formatted += "\n\n---\n\n## Images and Charts\n\n"
                for img_file in image_files:
                    rel_path = f"images/Session {session_num}_Handout_word/{img_file.name}"
                    formatted += f"![{img_file.stem}]({rel_path})\n\n"
        
        # Save improved markdown
        with open(output_md_path, 'w', encoding='utf-8') as f:
            f.write(formatted)
        
        print(f"✓ Saved improved markdown: {output_md_path}")
        print(f"  Final length: {len(formatted)} characters")
        return True
    else:
        print("✗ Failed to format with LLM")
        return False

def main():
    base_dir = Path("/Users/simonwang/Library/CloudStorage/OneDrive-HongKongBaptistUniversity/GTD/Areas/Teaching/Courses/MCCP 6020/PhDagentSpring2026")
    word_dir = base_dir / "Materials" / "word"
    md_dir = base_dir / "Materials" / "md"
    
    # Get API key
    api_key = get_poe_api_key()
    if not api_key:
        print("⚠ No Poe API key found")
        return
    
    # Process Session 1
    docx_file = word_dir / "Session 1_Handout_TEACHER version.docx"
    md_file = md_dir / "Session 1_Handout_TEACHER version_word.md"
    
    if not docx_file.exists():
        print(f"✗ Word file not found: {docx_file}")
        return
    
    improve_markdown_formatting(docx_file, md_file, api_key)
    print("\n✓ Formatting improvement complete!")

if __name__ == "__main__":
    main()
